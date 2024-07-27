import re
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import NavigableString

def get_alignment(element):
    if 'style' in element.attrs:
        if 'text-align: center' in element['style']:
            return WD_ALIGN_PARAGRAPH.CENTER
        elif 'text-align: right' in element['style']:
            return WD_ALIGN_PARAGRAPH.RIGHT
        elif 'text-align: justify' in element['style']:
            return WD_ALIGN_PARAGRAPH.JUSTIFY
    return WD_ALIGN_PARAGRAPH.LEFT

def get_font_size(element, default=12):
    font_size = element.find('span', style=lambda value: value and 'font-size' in value)
    if font_size:
        size = re.search(r'font-size:\s*(\d+)pt', font_size['style'])
        if size:
            return Pt(int(size.group(1)))
    return Pt(default)

def process_inline_elements(element, paragraph):
    for child in element.children:
        if isinstance(child, NavigableString):
            if child.strip():
                run = paragraph.add_run(child.string)
        elif child.name == 'br':
            paragraph.add_run().add_break()
        elif child.name in ['strong', 'b']:
            run = paragraph.add_run(child.get_text())
            run.bold = True
        elif child.name in ['em', 'i']:
            run = paragraph.add_run(child.get_text())
            run.italic = True
        elif child.name == 'u':
            run = paragraph.add_run(child.get_text())
            run.underline = True
        elif child.name == 'span':
            process_span(child, paragraph)
        elif child.name == 'sup':
            run = paragraph.add_run(child.get_text())
            run.font.superscript = True
        elif child.name == 'sub':
            run = paragraph.add_run(child.get_text())
            run.font.subscript = True
        else:
            process_inline_elements(child, paragraph)

def process_span(span, paragraph):
    run = paragraph.add_run(span.get_text())
    if 'style' in span.attrs:
        if 'text-decoration: underline' in span['style']:
            run.underline = True
        if 'text-decoration: line-through' in span['style']:
            run.font.strike = True
        if 'font-size' in span['style']:
            size = re.search(r'font-size:\s*(\d+)pt', span['style'])
            if size:
                run.font.size = Pt(int(size.group(1)))