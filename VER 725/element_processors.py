from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from style_utils import get_alignment, get_font_size, process_inline_elements
import re
from docx.shared import RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT


def process_heading(doc, element):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(element.get_text().strip())
    run.bold = True
    paragraph.alignment = get_alignment(element)
    run.font.size = get_font_size(element, default=16)

def process_paragraph(doc, element):
    if not element.get_text(strip=True):
            return
    paragraph = doc.add_paragraph()
    paragraph.alignment = get_alignment(element)
    process_inline_elements(element, paragraph)

def process_table(doc, element):
    # Get all rows from the table
    rows = element.find_all('tr')
    if not rows:
        return  # If there are no rows, return without creating a table

    # Get the number of columns from the first row
    first_row = rows[0]
    num_cols = len(first_row.find_all(['td', 'th']))

    # Create a table in the document
    table = doc.add_table(rows=len(rows), cols=num_cols)
    
    # Check if the table has the 'no-border' class
   
    if 'class' in element.attrs and 'no-border' in element['class']:
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        for row in table.rows:
            for cell in row.cells:
                for border in ['top', 'bottom', 'left', 'right']:
                    setattr(cell._element.get_or_add_tcPr(), f'w:{border}', None)
  

    # Process each row
    for i, row in enumerate(rows):
        cells = row.find_all(['td', 'th'])
        for j, cell in enumerate(cells):
            # Get the corresponding cell in the docx table
            table_cell = table.cell(i, j)
            
            # Clear default paragraph in cell
            table_cell.paragraphs[0].clear()

            # Process the content of the cell
            first_paragraph = True
            for child in cell.children:
                if child.name == 'p':
                  
                        if not first_paragraph:
                            paragraph = table_cell.add_paragraph()
                        else:
                            paragraph = table_cell.paragraphs[0]
                        first_paragraph = False
                        
                        # Set alignment for the paragraph
                        paragraph.alignment = get_alignment(child)
                        
                        # Remove space before and after paragraph
                        paragraph.paragraph_format.space_before = Pt(0)
                        paragraph.paragraph_format.space_after = Pt(0)
                        
                        # Process inline elements within the paragraph
                        process_inline_elements(child, paragraph)
                elif child.name == 'br':
                    paragraph = table_cell.add_paragraph()
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                elif child.string and child.string.strip():
                    # Handle direct text content
                    if first_paragraph:
                        paragraph = table_cell.add_paragraph()
                        first_paragraph = False
                    else:
                        paragraph = table_cell.paragraphs[-1]
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.add_run(child.string.strip())

            # Handle colspan
            if 'colspan' in cell.attrs:
                colspan = int(cell.attrs['colspan'])
                for k in range(1, colspan):
                    if j + k < num_cols:
                        table_cell.merge(table.cell(i, j + k))

            # Handle rowspan
            if 'rowspan' in cell.attrs:
                rowspan = int(cell.attrs['rowspan'])
                for k in range(1, rowspan):
                    if i + k < len(rows):
                        table_cell.merge(table.cell(i + k, j))

    # Handle table width
    if 'style' in element.attrs:
        style = element.attrs['style']
        width_match = re.search(r'width:\s*([\d.]+)%', style)
        if width_match:
            width_percent = float(width_match.group(1))
            table.width = Inches(6 * width_percent / 100)  # Assuming a default page width of 6 inches

    return table

def process_list(doc, element):
    if not element.get_text(strip=True):
            return
    for li in element.find_all('li'):
        paragraph = doc.add_paragraph(style='List Bullet')
        paragraph.alignment = get_alignment(li)
        process_inline_elements(li, paragraph)

def process_div(doc, element):
    paragraph = doc.add_paragraph()
    paragraph.alignment = get_alignment(element)
    process_inline_elements(element, paragraph)